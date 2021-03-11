import pandas as pd
from datetime import datetime
import matplotlib.pyplot as plt
import sys
import io

from docx import Document


plt.style.use('seaborn-darkgrid')
from matplotlib import rcParams
rcParams['font.family'] = 'monospace'
rcParams['font.sans-serif'] = ['Lucida Console']


def generate_report(city, date, dataset):
  def get_barchart(day_data):
    width = 0.35
    labels = day_data.query(query_day)['Product line'].unique()
    men_totals = []
    women_totals = []
    for category in labels:
      men_totals.append(day_data.query(f'`Product line` == "{category}" and Gender=="Male"')['Total'].sum())
      women_totals.append(day_data.query(f'`Product line` == "{category}" and Gender=="Female"')['Total'].sum())

    fig, ax = plt.subplots()
    ax.bar(labels, men_totals, width, label='Men')
    ax.bar(labels, women_totals, width, bottom=men_totals, label='Women')

    ax.set_ylabel('Revenue')
    ax.legend()

    ax.set_title('Revenue by category and gender')

    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    plt.close()
    return buf

  def compare(city, date, dataset):
    query = f"Date == '{date}'"
    ranking = dataset.query(query).groupby('City')[['Total']].sum()
    if(city == ranking.sort_values('Total', ascending=False).reset_index().iloc[0]['City']):
      return "Branch performed the best of all branches in terms of revenue!"
    elif (city == ranking.sort_values('Total', ascending=True).reset_index().iloc[0]['City']):
      return "Branch performed the worst of all branches in terms of revenue!"
    else:
      return False
  
  def compare_average(city, date, dataset, day_data):
    query = f"Date == '{date}'"
    global_average = dataset.query(query)['Total'].mean()
    average = day_data['Total'].mean()
    if (global_average < average):
      return (average, "higher than", global_average)
    elif (global_average > average):
      return (average, "lower than", global_average)
    else:
      return (average, "equal", global_average)

  query_day = f"City=='{city}' and Date=='{date}'"
  day_data = dataset.query(query_day)
  revenue = day_data['Total'].sum()

  document = Document()

  document.add_heading(f'Sales in {city} on {date}', 0)
  document.add_paragraph(f'The daily revenue was {revenue:0.2f}.')
  ranking = compare(city, date, dataset)
  if ranking:
    document.add_paragraph(ranking)
  verdict = compare_average(city, date, dataset, day_data)
  document.add_paragraph(f"The average profit per transaction was {verdict[0]:0.2f} and it was {verdict[1]} the global average ({verdict[2]:0.2f}).")

  document.add_picture(get_barchart(day_data))
  document.save(f'report_{city}_{date}.docx')



if __name__ == "__main__":
    data = pd.read_csv(sys.argv[1])

    city = sys.argv[2]
    date = sys.argv[3]
    generate_report(city, date, data)